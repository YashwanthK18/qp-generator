import os
import re
import random
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class QuestionPaperGenerator:
    def __init__(self, file_paths):
        self.file_paths = file_paths
        self.modules = {}
        self.output_path = os.path.join('output', 'new_question_paper.docx')
        
    def is_metadata(self, text):
        """Check if text is metadata (marks, CO's, PO's, etc.) and should be skipped"""
        # Single numbers or very short text
        if re.match(r'^\d+\s*$', text):
            return True
        
        # CO's, PO's, RBT indicators
        if re.match(r'^[1-5]\s*,?\s*[1-5]?\s*$', text):
            return True
            
        # Common metadata patterns
        metadata_patterns = [
            r'^Marks?$',
            r'^CO\'?s?$',
            r'^PO\'?s?$',
            r'^RBT$',
            r'^\d+\s*,\s*\d+$',
            r'^[1-5]$',
            r'Faculty In-charges',
            r'Course Coordinator',
            r'Programme Coordinator',
        ]
        
        return any(re.search(pattern, text, re.IGNORECASE) for pattern in metadata_patterns)
    
    def extract_questions_from_tables(self, doc):
        """Extract questions from document tables, filtering out metadata"""
        current_module = None
        questions_by_module = {}
        
        for table in doc.tables:
            current_question = None
            question_parts = []
            
            for row in table.rows:
                row_text = []
                
                # Get text from each cell, but only from likely question cells
                for i, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    
                    # Skip empty cells
                    if not text:
                        continue
                    
                    # Skip metadata columns (usually last 3-4 columns with numbers)
                    if i >= len(row.cells) - 4 and self.is_metadata(text):
                        continue
                    
                    row_text.append(text)
                
                if not row_text:
                    continue
                
                # Combine row text
                combined_text = ' '.join(row_text)
                
                # Check for module headers
                part_pattern = r'PART\s+([AB])\s*\(Module\s*[-–—]\s*([IVXivx]+)\)'
                part_match = re.search(part_pattern, combined_text, re.IGNORECASE)
                
                if part_match:
                    # Save previous question
                    if current_module and question_parts:
                        full_question = '\n'.join(question_parts)
                        if len(full_question) > 50:
                            questions_by_module[current_module].append(full_question)
                    
                    question_parts = []
                    module_num = part_match.group(2).upper()
                    
                    # Convert Roman to Arabic
                    roman_to_int = {'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5, 
                                   'VI': 6, 'VII': 7, 'VIII': 8, 'IX': 9, 'X': 10}
                    
                    module_number = roman_to_int.get(module_num, module_num)
                    current_module = f"Module {module_number}"
                    
                    if current_module not in questions_by_module:
                        questions_by_module[current_module] = []
                    
                    print(f"Found: {current_module}")
                    continue
                
                # Skip various header rows
                if any(keyword in combined_text for keyword in ['Q. No', 'Marks', 'OR']):
                    # Save question before OR
                    if combined_text.strip().upper() == 'OR' and question_parts:
                        full_question = '\n'.join(question_parts)
                        if len(full_question) > 50 and current_module:
                            questions_by_module[current_module].append(full_question)
                        question_parts = []
                    continue
                
                # Detect question start (e.g., "1. a)" or "2. a)")
                question_start = re.match(r'^(\d+)\.\s*([a-z])\)\s*(.+)', combined_text)
                
                if question_start:
                    # Save previous question
                    if question_parts and current_module:
                        full_question = '\n'.join(question_parts)
                        if len(full_question) > 50:
                            questions_by_module[current_module].append(full_question)
                    
                    # Start new question
                    question_num = question_start.group(1)
                    part_letter = question_start.group(2)
                    question_text = question_start.group(3).strip()
                    
                    # Clean out any remaining numeric metadata at the end
                    question_text = re.sub(r'\s+\d+\s*$', '', question_text)
                    
                    if len(question_text) > 10:  # Only if substantial
                        question_parts = [f"{part_letter}) {question_text}"]
                        current_question = question_num
                
                # Detect sub-parts (b), c), etc.)
                elif re.match(r'^([a-z])\)\s*(.+)', combined_text):
                    subpart_match = re.match(r'^([a-z])\)\s*(.+)', combined_text)
                    part_letter = subpart_match.group(1)
                    part_text = subpart_match.group(2).strip()
                    
                    # Clean out numeric metadata
                    part_text = re.sub(r'\s+\d+\s*$', '', part_text)
                    
                    if len(part_text) > 10:
                        question_parts.append(f"{part_letter}) {part_text}")
                
                # Continue existing question (if text is substantial and not metadata)
                elif current_question and not self.is_metadata(combined_text):
                    cleaned_text = combined_text.strip()
                    # Remove trailing numbers that are likely marks
                    cleaned_text = re.sub(r'\s+\d+\s*$', '', cleaned_text)
                    
                    if len(cleaned_text) > 20:
                        question_parts.append(cleaned_text)
            
            # Save last question from table
            if current_module and question_parts:
                full_question = '\n'.join(question_parts)
                if len(full_question) > 50:
                    questions_by_module[current_module].append(full_question)
                question_parts = []
        
        return questions_by_module
    
    def extract_questions_from_document(self, doc_path):
        """Extract modules and questions from a document"""
        try:
            doc = Document(doc_path)
            module_questions = self.extract_questions_from_tables(doc)
            
            # Merge into main modules dictionary
            for module, questions in module_questions.items():
                if module not in self.modules:
                    self.modules[module] = []
                self.modules[module].extend(questions)
            
            print(f"\nProcessed {os.path.basename(doc_path)}")
            for mod, questions in module_questions.items():
                print(f"  {mod}: {len(questions)} questions")
                for i, q in enumerate(questions, 1):
                    preview = q[:100].replace('\n', ' ')
                    print(f"    Q{i}: {preview}...")
            
            return True
            
        except Exception as e:
            print(f"Error processing {doc_path}: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
    
    def format_question_text(self, text):
        """Format question text properly"""
        lines = text.split('\n')
        formatted = []
        
        for line in lines:
            line = line.strip()
            if line:
                # Remove markdown bold markers
                line = re.sub(r'\*\*', '', line)
                formatted.append(line)
        
        return '\n'.join(formatted)
    
    def generate_paper(self):
        """Generate a new question paper"""
        try:
            # Extract questions from all documents
            for file_path in self.file_paths:
                print(f"\n{'='*70}")
                print(f"Processing: {os.path.basename(file_path)}")
                print('='*70)
                self.extract_questions_from_document(file_path)
            
            if not self.modules:
                return {
                    'success': False,
                    'error': 'No modules found in the uploaded documents'
                }
            
            # Filter out modules with no questions
            self.modules = {k: v for k, v in self.modules.items() if v}
            
            if not self.modules:
                return {
                    'success': False,
                    'error': 'No questions found in any modules'
                }
            
            # Create new document
            new_doc = Document()
            
            # Add header
            header = new_doc.add_heading('BANGALORE INSTITUTE OF TECHNOLOGY', 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            dept = new_doc.add_paragraph('DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING')
            dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dept.runs[0].bold = True
            
            title = new_doc.add_paragraph('Generated Question Paper - 2024-25')
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].bold = True
            
            new_doc.add_paragraph()
            
            # Course info
            info = new_doc.add_paragraph()
            info.add_run('COURSE: ').bold = True
            info.add_run('Theory of Computation (BCS503)     ')
            info.add_run('MAX MARKS: ').bold = True
            info.add_run('25')
            
            new_doc.add_paragraph()
            
            # Note
            note = new_doc.add_paragraph()
            note.add_run('Note: ').bold = True
            note.add_run('Answer any one full question from each PART')
            
            new_doc.add_paragraph()
            
            # Sort modules
            sorted_modules = sorted(
                self.modules.keys(), 
                key=lambda x: int(re.search(r'\d+', x).group()) if re.search(r'\d+', x) else 0
            )
            
            question_num = 1
            
            for module_name in sorted_modules:
                questions = self.modules[module_name]
                
                if not questions:
                    continue
                
                # Module heading
                heading = new_doc.add_heading(f'PART A ({module_name})', level=1)
                
                # First question
                selected_q1 = random.choice(questions)
                formatted_q1 = self.format_question_text(selected_q1)
                
                q_para = new_doc.add_paragraph()
                q_para.add_run(f'{question_num}. ').bold = True
                q_para.add_run(formatted_q1)
                
                new_doc.add_paragraph()
                
                # OR
                or_para = new_doc.add_paragraph('OR')
                or_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                or_para.runs[0].bold = True
                
                new_doc.add_paragraph()
                
                # Second question
                question_num += 1
                remaining = [q for q in questions if q != selected_q1]
                selected_q2 = random.choice(remaining) if remaining else selected_q1
                formatted_q2 = self.format_question_text(selected_q2)
                
                q_para = new_doc.add_paragraph()
                q_para.add_run(f'{question_num}. ').bold = True
                q_para.add_run(formatted_q2)
                
                new_doc.add_paragraph()
                question_num += 1
            
            # Save
            os.makedirs('output', exist_ok=True)
            new_doc.save(self.output_path)
            
            print(f"\n{'='*70}")
            print(f"SUCCESS! Generated paper from {len(self.modules)} modules")
            print('='*70)
            
            return {
                'success': True,
                'modules_found': len(self.modules),
                'questions_extracted': sum(len(q) for q in self.modules.values()),
                'output_path': self.output_path
            }
            
        except Exception as e:
            import traceback
            print(f"\nERROR: {str(e)}")
            traceback.print_exc()
            return {
                'success': False,
                'error': str(e)
            }

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        files = sys.argv[1:]
        generator = QuestionPaperGenerator(files)
        result = generator.generate_paper()
        print("\nFinal Result:")
        print(result)
    else:
        print("Usage: python question_generator.py <file1.docx> <file2.docx> ...")